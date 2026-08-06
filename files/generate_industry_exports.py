#!/usr/bin/env python3
"""Generate PDF and DOCX for the industry resume (content mirrors industry.html)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Flowable,
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

OUT_DIR = Path(__file__).resolve().parent
PDF_PATH = OUT_DIR / "Md_Ashikur_Rahman_Industry_Resume.pdf"
DOCX_PATH = OUT_DIR / "Md_Ashikur_Rahman_Industry_Resume.docx"

NAVY = HexColor("#00274c")
MUTED = HexColor("#5f6b75")
LINE = HexColor("#e4eaee")
BODY = HexColor("#222222")

NAVY_RGB = RGBColor(0x00, 0x27, 0x4C)
BODY_RGB = RGBColor(0x22, 0x22, 0x22)
MUTED_RGB = RGBColor(0x5F, 0x6B, 0x75)
MAROON_RGB = RGBColor(0x9B, 0x1B, 0x30)

HEADLINE = "Lead AI Engineer | Multimodal AI, Generative AI & Applied Computer Vision"

SUMMARY = (
    "Lead AI Engineer specializing in multimodal AI, vision-language systems, generative AI, "
    "and applied computer vision. Currently lead a multidisciplinary team of 15+ engineers and "
    "researchers and deliver enterprise AI products used by 200+ brands. Architected "
    "computer-vision systems that have processed more than 4.5 million images globally."
)

EXPERIENCE = [
    (
        "Lead AI Engineer, The KOW Company",
        "Jan 2023-Present",
        [
            "Lead a multidisciplinary team of 15+ ML engineers, software engineers, and junior researchers, driving technical strategy, system architecture, engineering execution, quality standards, and cross-functional delivery.",
            "Lead the design, development, and deployment of scalable AI systems, translating product requirements into robust machine learning solutions.",
            "Direct applied research on AI hallucination, prompt safety, visual grounding, and video understanding, operationalizing research into production-ready models and evaluation frameworks.",
            "Develop and deploy deep learning–based computer vision and image-processing systems; publish reproducible code, datasets, and model checkpoints on GitHub and Hugging Face.",
        ],
    ),
    (
        "Senior Machine Learning Engineer, The KOW Company",
        "Jul 2021-Dec 2022",
        [
            "Improved object detection and segmentation performance by 20-35% across internal evaluation benchmarks; the resulting models were later deployed in Retouched.ai.",
            "Led 6+ client ML engagements from requirements gathering through production delivery; built offline evaluation pipelines and A/B testing workflows to validate model quality, inference performance, and business and operational outcomes.",
        ],
    ),
    (
        "Machine Learning Engineer, The KOW Company",
        "Jul 2020-Jun 2021",
        [
            "Built deep learning models for object recognition, image segmentation, and background-removal workflows; developed scalable preprocessing, training, and A/B testing pipelines.",
        ],
    ),
    (
        "Senior Software Engineer, Smart Technologies (BD) Ltd",
        "Sep 2016-Dec 2019",
        [
            "Led .NET and SQL Server supply-chain systems on a 4 TB database, reducing report generation from 20 minutes to 40-54 seconds.",
            "Achieved 70-75% process automation and 99.9% synchronization success for offline-capable enterprise workflows.",
        ],
    ),
    (
        "Software Engineer, Proggasoft",
        "Mar 2015-Aug 2016",
        [
            "Developed ASP.NET MVC features, backend services, and database integrations for DevSkill.com.",
        ],
    ),
]

PROJECTS = [
    (
        '<link href="https://retouched.ai" color="#9b1b30"><u>Retouched.ai</u></link> - Object Detection and Segmentation',
        "Production",
        [
            "Developed salient-object segmentation for background removal; improved segmentation quality by 17% on internal benchmarks, reduced processing time by 30%, enabled uploads of up to 257 MB, and achieved a 2.27-second average processing time across standard workloads.",
            "Scaled Retouched.ai to process 4.5M+ images globally for hundreds of customers using PyTorch, U²-Net-inspired salient-object segmentation, FastAPI, and Google Cloud Platform.",
        ],
    ),
    (
        '<link href="https://omnimage.ai" color="#9b1b30"><u>Omnimage.ai</u></link> - AI Image and Video Generation',
        "Production",
        [
            "Co-designed and launched image- and video-generation APIs used by 200+ brands for creative and product-image workflows.",
            "Engineered workflows for reference-image conditioning, asynchronous processing, prompt classification, intent routing, and automated model selection.",
        ],
    ),
    (
        "The Fitting Room - Cross-Brand Virtual Try-On Platform (In development)",
        "",
        [
            "Conceived and co-led a unified cross-brand virtual try-on platform supporting products from 170+ brands.",
            "Architected a Dockerized FastAPI/Nginx backend using SQL Server, Google Cloud Storage, Redis, recommendation services, and 2D virtual try-on pipelines.",
        ],
    ),
    (
        "Enterprise AI Catalog Audit and Image Quality Assurance Platform (In development)",
        "",
        [
            "Leading the development of an AI catalog-audit and image quality assurance platform for a major US retail client.",
            "Automating image QA, catalog validation, metadata checks, and content-health monitoring across DAM and CMS workflows using NLP, PyTorch, and FastAPI.",
        ],
    ),
    (
        "CogniX - Proprietary Multimodal Content Creation Platform (Active R&D)",
        "",
        [
            "Leading early-stage R&D for multimodal product-visualization and editing workflows.",
            "Researching and prototyping multi-reference fusion and editing workflows for garment replacement, scene composition, product visualization, and style transfer using PyTorch, Diffusers, ComfyUI, and parameter-efficient fine-tuning (LoRA, QLoRA).",
        ],
    ),
]

PEER_REVIEWED = [
    '[P1] Abdullah Ibne Hanif Arean, Niamul Hassan Samin, Md Arifur Rahman, Renu Akter Sweety, '
    'Juena Ahmed Noshin, <b>Md Ashikur Rahman</b>*. '
    '&ldquo;Stroke-Level Connectivity Verification: Grounding Vision-Language Models Against '
    'Topology Hallucination in Diagram Understanding.&rdquo; '
    '<i>International Conference on Document Analysis and Recognition (ICDAR)</i>, 2026. '
    'Accepted for oral presentation. *Corresponding author. '
    '<link href="https://github.com/tkcl-research/LogicBench1k" color="#9b1b30"><u>[Code]</u></link>',
    '[P2] Juena Noshin, Kawser Irom Rushee, Mohammad Rabiul Islam, Sifat Rahman Ahona, '
    '<b>Md Ashikur Rahman</b>. '
    '&ldquo;UCAR: Uncertainty-Calibrated Adaptive Retrieval for Hallucination Reduction in '
    'Medical Vision-Language Models.&rdquo; '
    '<i>International Conference on Computing Advancements (ICCA)</i>, 2026.',
]

ADDITIONAL_PUBLICATION = [
    '[A1] <b>Md Ashikur Rahman</b>, Md Arifur Rahman, Juena Ahmed Noshin. '
    '&ldquo;Automated Detection of Diabetic Retinopathy Using Deep Residual Learning.&rdquo; '
    '<i>International Journal of Computer Applications</i>, 2020.',
]

PREPRINTS = [
    '[M1] <b>Md Ashikur Rahman</b>, Md Arifur Rahman, Niamul Hassan Samin, Khandaker Rifah Tasnia, '
    'Sifat Rahman Ahona, Juena Ahmed Noshin. '
    '&ldquo;Beyond Aggregate Risk: Role-Stratified Conformal Risk Control for LLM Tool Calls.&rdquo; '
    'Manuscript under review, 2026.',
    '[M2] <b>Md Ashikur Rahman</b>, Juena Ahmed Noshin, Niamul Hassan Samin, Abdullah Ibne Hanif Arean, '
    'Md Hasibul Amin, Md Arifur Rahman. '
    '&ldquo;When Detector-Based Grounding Metrics Measure Vocabulary: A Cautionary Audit of Entity '
    'Claims in Video-QA Reasoning Traces.&rdquo; '
    'Manuscript under review, 2026.',
    '[M3] <b>Md Ashikur Rahman</b>, Md Arifur Rahman, Nusrat Jahan Trisna, Juena Ahmed Noshin. '
    '&ldquo;Decoding Harm: Do Reasoning Models Resist Math-Encoded Jailbreaks?&rdquo; '
    'Manuscript under review, 2026.',
    '[M4] Niamul Hassan Samin, Abdullah Ibne Hanif Arean, Md Arifur Rahman, Md Hasibul Amin, '
    'Renu Akter Suity, Juena Ahmed Noshin, <b>Md Ashikur Rahman</b>. '
    '&ldquo;Residual Stream Rebalancing: Training-Free Hallucination Mitigation in '
    'Vision-Language Models.&rdquo; '
    'Manuscript under review, 2026.',
]

SKILLS = [
    ("Multimodal and Vision-Language AI:", "Python, PyTorch, Vision-Language Models, Visual Grounding, Hallucination Evaluation, Multimodal Learning"),
    ("Generative AI and LLMs:", "Diffusion Models, Large Language Models (including Llama and Qwen), Prompt Safety, Parameter-Efficient Fine-Tuning (LoRA, QLoRA)"),
    ("Computer Vision:", "Object Detection, Segmentation, Pose Estimation, Image Quality Assurance"),
    ("3D Vision:", "Structure from Motion, Multi-View Stereo, COLMAP, Open3D, Neural Radiance Fields (NeRF), 3D Gaussian Splatting"),
    ("Production ML and MLOps:", "Model Serving, Deployment, Offline and Online Evaluation, A/B Testing, Dataset Design"),
    ("Backend and Cloud:", "FastAPI, Docker, Nginx, Redis, Google Cloud Platform, Google Cloud Storage, SQL Server"),
]

LOCATION_LINE = "Dhaka, Bangladesh | Open to Relocation and Remote Opportunities"

CONTACT_PARTS = [
    ("link", "mdashikur.rafi@gmail.com", "mailto:mdashikur.rafi@gmail.com"),
    ("text", " | +880 1675 964 080 | "),
    ("link", "Portfolio", "https://ashikrafi.github.io/"),
    ("text", " | "),
    ("link", "GitHub", "https://github.com/ashikrafi"),
    ("text", " | "),
    ("link", "LinkedIn", "https://www.linkedin.com/in/mdashikrah/"),
    ("text", " | "),
    ("link", "Google Scholar", "https://scholar.google.com/citations?user=Htgw_vEAAAAJ&hl=en"),
]


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ResumeName",
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=20,
            alignment=TA_CENTER,
            textColor=NAVY,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeHeadline",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            alignment=TA_CENTER,
            textColor=NAVY,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeContact",
            fontName="Helvetica",
            fontSize=8,
            leading=10.5,
            alignment=TA_CENTER,
            textColor=NAVY,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionTitle",
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=NAVY,
            spaceBefore=5,
            spaceAfter=2,
            leftIndent=0,
            firstLineIndent=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Summary",
            fontName="Helvetica",
            fontSize=9,
            leading=11.2,
            alignment=TA_LEFT,
            textColor=BODY,
            spaceAfter=2,
            leftIndent=0,
            firstLineIndent=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EntryTitle",
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=11.5,
            textColor=BODY,
            leftIndent=0,
            firstLineIndent=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EntryDates",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11.5,
            alignment=TA_RIGHT,
            textColor=MUTED,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeBullet",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            leftIndent=0,
            firstLineIndent=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Degree",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            textColor=NAVY,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EduMeta",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PubItem",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            leftIndent=0,
            firstLineIndent=0,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SkillLine",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            spaceAfter=1,
        )
    )
    return styles


class EntryHeader(Flowable):
    """Title at frame left edge; date right-aligned on the same line (no Table overhang)."""

    def __init__(self, title: str, dates: str, styles, date_width=1.55 * inch):
        Flowable.__init__(self)
        self._title = Paragraph(
            title if "<link" in title else _escape(title),
            styles["EntryTitle"],
        )
        self._dates = Paragraph(_escape(dates), styles["EntryDates"]) if dates else None
        self._date_width = date_width
        self._title_h = 0
        self._dates_h = 0

    def wrap(self, availWidth, availHeight):
        self.width = availWidth
        gap = 6 if self._dates else 0
        title_w = availWidth - (self._date_width + gap if self._dates else 0)
        _, self._title_h = self._title.wrap(title_w, availHeight)
        self._dates_h = 0
        if self._dates:
            _, self._dates_h = self._dates.wrap(self._date_width, availHeight)
        self.height = max(self._title_h, self._dates_h) + 1
        return self.width, self.height

    def draw(self):
        y_title = self.height - self._title_h
        self._title.drawOn(self.canv, 0, y_title)
        if self._dates:
            y_dates = self.height - self._dates_h
            self._dates.drawOn(self.canv, self.width - self._date_width, y_dates)


def entry_header_table(title: str, dates: str, styles) -> Flowable:
    return EntryHeader(title, dates, styles)


def bullets(items: list[str], styles) -> list:
    return [Paragraph(f"• {_escape(item)}", styles["ResumeBullet"]) for item in items]


def section_rule() -> HRFlowable:
    return HRFlowable(width="100%", thickness=1, color=LINE, spaceBefore=0, spaceAfter=3)


def contact_pdf_html() -> str:
    parts = []
    for item in CONTACT_PARTS:
        if item[0] == "text":
            parts.append(_escape(item[1]))
        else:
            parts.append(f'<link href="{item[2]}" color="#9b1b30"><u>{_escape(item[1])}</u></link>')
    return "".join(parts)


def build_pdf():
    styles = build_pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.4 * inch,
        bottomMargin=0.4 * inch,
        title="Md Ashikur Rahman | Industry Resume",
        author="Md Ashikur Rahman",
    )

    story = []
    story.append(Paragraph("MD ASHIKUR RAHMAN", styles["ResumeName"]))
    story.append(Paragraph(_escape(HEADLINE), styles["ResumeHeadline"]))
    story.append(Paragraph(_escape(LOCATION_LINE), styles["ResumeContact"]))
    story.append(Paragraph(contact_pdf_html(), styles["ResumeContact"]))

    story.append(Paragraph("Summary", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(Paragraph(_escape(SUMMARY), styles["Summary"]))

    story.append(Paragraph("Experience", styles["SectionTitle"]))
    story.append(section_rule())
    for title, dates, items in EXPERIENCE:
        block = [entry_header_table(title, dates, styles)]
        block.extend(bullets(items, styles))
        block.append(Spacer(1, 3))
        story.append(KeepTogether(block))

    story.append(Paragraph("Key Projects", styles["SectionTitle"]))
    story.append(section_rule())
    for title, status, items in PROJECTS:
        block = [entry_header_table(title, status, styles)]
        block.extend(bullets(items, styles))
        block.append(Spacer(1, 3))
        story.append(KeepTogether(block))

    story.append(PageBreak())

    pub_block = [
        Paragraph("Peer-Reviewed Publications", styles["SectionTitle"]),
        section_rule(),
    ]
    for pub in PEER_REVIEWED:
        pub_block.append(Paragraph(pub, styles["PubItem"]))
    story.append(KeepTogether(pub_block))

    preprint_block = [
        Paragraph("Preprints and Manuscripts Under Review", styles["SectionTitle"]),
        section_rule(),
    ]
    for text in PREPRINTS:
        preprint_block.append(Paragraph(text, styles["PubItem"]))
    story.append(KeepTogether(preprint_block))

    addl_block = [
        Paragraph("Additional Publication", styles["SectionTitle"]),
        section_rule(),
    ]
    for pub in ADDITIONAL_PUBLICATION:
        addl_block.append(Paragraph(pub, styles["PubItem"]))
    story.append(KeepTogether(addl_block))

    story.append(Paragraph("Education", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(Paragraph("B.Sc. in Computer Science and Engineering", styles["Degree"]))
    story.append(
        Paragraph(
            "American International University-Bangladesh | 2011-2015",
            styles["EntryTitle"],
        )
    )
    story.append(Paragraph("CGPA: 3.87/4.00; WES equivalent: 3.94/4.00", styles["EduMeta"]))
    story.append(
        Paragraph(
            "Magna Cum Laude; Top 3%; Merit Scholarship and Tuition Fee Waiver",
            styles["EduMeta"],
        )
    )

    story.append(Paragraph("Awards", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(
        Paragraph(
            "• Champion, BASIS National ICT Awards 2020 (Retouched.ai); Finalist, APICTA 2021",
            styles["PubItem"],
        )
    )

    story.append(Paragraph("Languages", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(
        Paragraph(
            "• Bengali (Native); English (Professional working proficiency)",
            styles["PubItem"],
        )
    )

    story.append(Paragraph("Skills", styles["SectionTitle"]))
    story.append(section_rule())
    for label, text in SKILLS:
        story.append(
            Paragraph(f"<b>{_escape(label)}</b> {_escape(text)}", styles["SkillLine"])
        )

    doc.build(story)
    print(f"Wrote {PDF_PATH}")


def set_run_font(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E4EAEE")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2, line=1.1)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, color=NAVY_RGB)
    add_bottom_border(p)
    return p


def plain_title(title: str) -> str:
    """Strip reportlab link/markup for plain-text consumers (e.g. DOCX)."""
    import re

    text = re.sub(r'<link[^>]*>', "", title)
    for tag in ("</link>", "<u>", "</u>", "<b>", "</b>", "<i>", "</i>"):
        text = text.replace(tag, "")
    return text


def add_entry_header(doc, title, dates=""):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=1, line=1.1)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    display_title = plain_title(title)
    if dates:
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
        run = p.add_run(display_title)
        set_run_font(run, size=10, bold=True, color=BODY_RGB)
        p.add_run("\t")
        date_run = p.add_run(dates)
        set_run_font(date_run, size=9, bold=False, color=MUTED_RGB)
    else:
        run = p.add_run(display_title)
        set_run_font(run, size=10, bold=True, color=BODY_RGB)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=1, line=1.15)
    # Keep bullets indented from the shared left edge (same as section titles / job titles)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run("• " + text)
    set_run_font(run, size=9.5, bold=False, color=BODY_RGB)
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "9B1B30")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, before=0, after=2, line=1.0)
    run = name.add_run("MD ASHIKUR RAHMAN")
    set_run_font(run, size=18, bold=True, color=NAVY_RGB)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(headline, before=0, after=2, line=1.1)
    run = headline.add_run(HEADLINE)
    set_run_font(run, size=10, bold=True, color=NAVY_RGB)

    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(location, before=0, after=1, line=1.1)
    run = location.add_run(LOCATION_LINE)
    set_run_font(run, size=9, bold=False, color=NAVY_RGB)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, before=0, after=6, line=1.1)
    for item in CONTACT_PARTS:
        if item[0] == "text":
            run = contact.add_run(item[1])
            set_run_font(run, size=9, bold=False, color=NAVY_RGB)
        else:
            add_hyperlink(contact, item[1], item[2])

    add_section_heading(doc, "Summary")
    summary = doc.add_paragraph()
    set_paragraph_spacing(summary, before=2, after=4, line=1.15)
    summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
    summary.paragraph_format.left_indent = Inches(0)
    summary.paragraph_format.first_line_indent = Inches(0)
    run = summary.add_run(SUMMARY)
    set_run_font(run, size=9.5, color=BODY_RGB)

    add_section_heading(doc, "Experience")
    for title, dates, items in EXPERIENCE:
        add_entry_header(doc, title, dates)
        for item in items:
            add_bullet(doc, item)

    add_section_heading(doc, "Key Projects")
    for title, status, items in PROJECTS:
        add_entry_header(doc, title, status)
        for item in items:
            add_bullet(doc, item)

    doc.add_page_break()

    add_section_heading(doc, "Peer-Reviewed Publications")
    for pub in PEER_REVIEWED:
        plain = (
            pub.replace("<i>", "")
            .replace("</i>", "")
            .replace("<b>", "")
            .replace("</b>", "")
            .replace("&ldquo;", '"')
            .replace("&rdquo;", '"')
            .replace(
                '<link href="https://github.com/tkcl-research/LogicBench1k" color="#9b1b30"><u>[Code]</u></link>',
                "[Code]",
            )
        )
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(plain)
        set_run_font(run, size=9.5, bold=False, color=BODY_RGB)

    add_section_heading(doc, "Preprints and Manuscripts Under Review")
    for text in PREPRINTS:
        plain = (
            text.replace("<b>", "")
            .replace("</b>", "")
            .replace("&ldquo;", '"')
            .replace("&rdquo;", '"')
        )
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(plain)
        set_run_font(run, size=9.5, bold=False, color=BODY_RGB)

    add_section_heading(doc, "Additional Publication")
    for pub in ADDITIONAL_PUBLICATION:
        plain = (
            pub.replace("<i>", "")
            .replace("</i>", "")
            .replace("&ldquo;", '"')
            .replace("&rdquo;", '"')
        )
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(plain)
        set_run_font(run, size=9.5, bold=False, color=BODY_RGB)

    add_section_heading(doc, "Education")
    degree = doc.add_paragraph()
    set_paragraph_spacing(degree, before=2, after=1, line=1.1)
    run = degree.add_run("B.Sc. in Computer Science and Engineering")
    set_run_font(run, size=10, bold=True, color=NAVY_RGB)
    add_entry_header(doc, "American International University-Bangladesh | 2011-2015")
    for line in (
        "CGPA: 3.87/4.00; WES equivalent: 3.94/4.00",
        "Magna Cum Laude; Top 3%; Merit Scholarship and Tuition Fee Waiver",
    ):
        meta = doc.add_paragraph()
        set_paragraph_spacing(meta, before=0, after=1, line=1.15)
        run = meta.add_run(line)
        set_run_font(run, size=9.5, color=BODY_RGB)

    add_section_heading(doc, "Awards")
    add_bullet(
        doc,
        "Champion, BASIS National ICT Awards 2020 (Retouched.ai); Finalist, APICTA 2021",
    )

    add_section_heading(doc, "Languages")
    add_bullet(doc, "Bengali (Native); English (Professional working proficiency)")

    add_section_heading(doc, "Skills")
    for label, text in SKILLS:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=1, line=1.15)
        run = p.add_run(label + " ")
        set_run_font(run, size=9.5, bold=True, color=BODY_RGB)
        run2 = p.add_run(text)
        set_run_font(run2, size=9.5, bold=False, color=BODY_RGB)

    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build_pdf()
    build_docx()
