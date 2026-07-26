#!/usr/bin/env python3
"""Generate PDF and DOCX for the industry resume (content mirrors industry.html)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT_DIR = Path(__file__).resolve().parent
PDF_PATH = OUT_DIR / "Md_Ashikur_Rahman_Industry_Resume.pdf"
DOCX_PATH = OUT_DIR / "Md_Ashikur_Rahman_Industry_Resume.docx"

NAVY = HexColor("#00274c")
MUTED = HexColor("#5f6b75")
LINE = HexColor("#e4eaee")
BODY = HexColor("#222222")

NAVY_RGB = RGBColor(0x00, 0x27, 0x4C)
BODY_RGB = RGBColor(0x22, 0x22, 0x22)
MUTED_RGB = RGBColor(0x5F, 0x6B, 0x75)
MAROON_RGB = RGBColor(0x9B, 0x1B, 0x30)

HEADLINE = "Lead AI Engineer | Multimodal AI, Generative AI, and Production Computer Vision"

SUMMARY = (
    "Lead AI Engineer specializing in multimodal AI, vision-language systems, generative AI, "
    "and production computer vision. Lead 15+ engineers and researchers across applied research, "
    "product development, and enterprise delivery. Built and scaled production AI systems "
    "supporting 200+ brands, including computer-vision platforms that have processed 4.5M+ images "
    "globally."
)

EXPERIENCE = [
    (
        "Lead AI Engineer, The KOW Company",
        "Jan 2023-Present",
        [
            "Lead a multidisciplinary team of 15+ ML engineers, software engineers, and researchers across architecture, planning, code review, and delivery; maintain 90-95% on-time delivery across AI engineering projects.",
            "Oversee development of production AI systems across virtual try-on, generative media, 3D reconstruction, catalog audit, and audio QA.",
            "Direct applied research on AI hallucination, prompt safety, visual grounding, and video understanding.",
            "Develop and scale production segmentation and image-processing systems, including Retouched.ai; publish reproducible research code, datasets, and model artifacts on GitHub and Hugging Face.",
        ],
    ),
    (
        "Senior Machine Learning Engineer, The KOW Company",
        "Jul 2021-Dec 2022",
        [
            "Improved object detection and segmentation performance by 20-35% across internal evaluation benchmarks; the resulting models were later deployed through Retouched.ai.",
            "Led 6+ client ML engagements from business requirements to technical delivery; built offline evaluation pipelines and production A/B testing workflows to validate model quality, inference performance, and production outcomes.",
        ],
    ),
    (
        "Machine Learning Engineer, The KOW Company",
        "Jul 2020-Jun 2021",
        [
            "Built deep learning models for production object recognition, image segmentation, and background-removal workflows; developed scalable preprocessing, training, and A/B testing pipelines.",
        ],
    ),
    (
        "Senior Software Engineer, Smart Technologies (BD) Ltd",
        "Sep 2016-Dec 2019",
        [
            "Led .NET and SQL Server supply-chain systems on a 4 TB database, reducing report generation from 20 minutes to 40-54 seconds.",
            "Achieved 70-75% process automation and 99.9% synchronization success for offline-capable enterprise workflows.",
        ],
    ),
    (
        "Software Engineer, Proggasoft",
        "Mar 2015-Aug 2016",
        [
            "Developed ASP.NET MVC features, backend services, and database integrations for DevSkill.com.",
        ],
    ),
]

PROJECTS = [
    (
        "Retouched.ai - Object Detection and Segmentation",
        "Production",
        [
            "Developed production salient-object segmentation for background removal; improved segmentation quality by 17% on internal benchmarks, reduced processing time by 30%, supported uploads up to 257 MB, and achieved a 2.27-second average processing time across standard production workloads.",
            "Scaled Retouched.ai to process 4.5M+ images globally for hundreds of customers using PyTorch, U²-Net-inspired salient-object segmentation, FastAPI, and Google Cloud Platform.",
        ],
    ),
    (
        "Omnimage.ai - AI Image and Video Generation",
        "Production",
        [
            "Co-designed and launched production image- and video-generation APIs used across 200+ brands for creative and product-image workflows.",
            "Built workflows for reference-image conditioning, asynchronous processing, prompt classification, intent routing, and automated model selection.",
        ],
    ),
    (
        "The Fitting Room - Cross-Brand Virtual Try-On Platform",
        "In Development",
        [
            "Conceived and co-led a unified cross-brand virtual try-on platform covering 170+ brands.",
            "Architected a Dockerized FastAPI/Nginx backend using SQL Server, Google Cloud Storage, Redis, recommendation services, and 2D virtual try-on pipelines.",
        ],
    ),
    (
        "Enterprise AI Catalog Audit and Image Quality Assurance Platform",
        "In Development",
        [
            "Lead development of an AI catalog-audit and image quality assurance platform for a major US retail client.",
            "Automate image QA, catalog validation, metadata checks, and content-health monitoring across DAM and CMS workflows using computer vision, NLP, PyTorch, and FastAPI.",
        ],
    ),
    (
        "CogniX - Proprietary Multimodal Image Generation and Editing",
        "Active R&D",
        [
            "Leading early-stage R&D for a proprietary multimodal image-generation and editing system for product-visualization workflows.",
            "Researching and prototyping multi-reference fusion and prompt-guided image-editing workflows for garment replacement, scene composition, product visualization, and style transfer using PyTorch, Diffusers, ComfyUI, and parameter-efficient fine-tuning (LoRA, QLoRA).",
        ],
    ),
]

PEER_REVIEWED = [
    "Stroke-Level Connectivity Verification: Grounding Vision-Language Models Against Topology Hallucination in Diagram Understanding. Accepted for oral presentation at ICDAR 2026; Corresponding Author.",
    "Automated Detection of Diabetic Retinopathy using Deep Residual Learning. IJCA, 2020.",
]

PREPRINTS = [
    (
        "Step-Level Visual Grounding Faithfulness Predicts Out-of-Distribution Generalization in Long-Horizon Vision-Language Models. arXiv preprint, 2026. Under review.",
        "arxiv.org/pdf/2603.06828",
        "https://arxiv.org/pdf/2603.06828",
    ),
    (
        "Beyond Dominant Patches: Spatial Credit Redistribution for Grounded Vision-Language Models. arXiv preprint, 2026. Under review.",
        "arxiv.org/pdf/2602.22469",
        "https://arxiv.org/pdf/2602.22469",
    ),
]

SKILLS = [
    ("Multimodal and Vision-Language AI:", "Python, PyTorch, Vision-Language Models, Visual Grounding, Hallucination Evaluation, Multimodal Learning"),
    ("Generative AI and LLMs:", "Diffusion Models, Large Language Models (Llama, Qwen), Prompt Safety, Parameter-Efficient Fine-Tuning (LoRA, QLoRA), Prompt-Guided Image Editing"),
    ("Computer Vision:", "Object Detection, Segmentation, Pose Estimation, Image Quality Assurance"),
    ("3D Vision:", "Structure from Motion, Multi-View Stereo, COLMAP, Open3D, Neural Radiance Fields (NeRF), 3D Gaussian Splatting"),
    ("Production ML and MLOps:", "Model Serving, Production Deployment, Offline Evaluation, Production A/B Testing, Model Evaluation, Dataset Design"),
    ("Backend and Cloud:", "FastAPI, Docker, Nginx, Redis, Google Cloud Platform, Google Cloud Storage, SQL Server"),
]

LOCATION_LINE = "Dhaka, Bangladesh | Open to Relocation and Remote Opportunities"

CONTACT_PARTS = [
    ("link", "mdashikur.rafi@gmail.com", "mailto:mdashikur.rafi@gmail.com"),
    ("text", " | +880 1675 964 080 | "),
    ("link", "Portfolio", "https://ashikrafi.github.io/"),
    ("text", " | "),
    ("link", "GitHub", "https://github.com/ashikrafi"),
    ("text", " | "),
    ("link", "LinkedIn", "https://www.linkedin.com/in/mdashikrah/"),
    ("text", " | "),
    ("link", "Google Scholar", "https://scholar.google.com/citations?user=Htgw_vEAAAAJ&hl=en"),
]


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ResumeName",
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=20,
            alignment=TA_CENTER,
            textColor=NAVY,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeHeadline",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            alignment=TA_CENTER,
            textColor=NAVY,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeContact",
            fontName="Helvetica",
            fontSize=8,
            leading=10.5,
            alignment=TA_CENTER,
            textColor=NAVY,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionTitle",
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=NAVY,
            spaceBefore=5,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Summary",
            fontName="Helvetica",
            fontSize=9,
            leading=11.2,
            alignment=TA_LEFT,
            textColor=BODY,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EntryTitle",
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=11.5,
            textColor=BODY,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EntryDates",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11.5,
            alignment=TA_RIGHT,
            textColor=MUTED,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeBullet",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            leftIndent=10,
            bulletIndent=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Degree",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            textColor=NAVY,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EduMeta",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PubItem",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            leftIndent=10,
            bulletIndent=0,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SkillLine",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=BODY,
            spaceAfter=1,
        )
    )
    return styles


def entry_header_table(title: str, dates: str, styles) -> Table:
    if dates:
        data = [
            [
                Paragraph(_escape(title), styles["EntryTitle"]),
                Paragraph(_escape(dates), styles["EntryDates"]),
            ]
        ]
        col_widths = [5.2 * inch, 2.3 * inch]
    else:
        data = [[Paragraph(_escape(title), styles["EntryTitle"])]]
        col_widths = [7.5 * inch]
    table = Table(data, colWidths=col_widths)
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    return table


def bullets(items: list[str], styles) -> list:
    return [Paragraph(f"• {_escape(item)}", styles["ResumeBullet"]) for item in items]


def section_rule() -> HRFlowable:
    return HRFlowable(width="100%", thickness=1, color=LINE, spaceBefore=0, spaceAfter=3)


def contact_pdf_html() -> str:
    parts = []
    for item in CONTACT_PARTS:
        if item[0] == "text":
            parts.append(_escape(item[1]))
        else:
            parts.append(f'<link href="{item[2]}" color="#9b1b30"><u>{_escape(item[1])}</u></link>')
    return "".join(parts)


def build_pdf():
    styles = build_pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.4 * inch,
        bottomMargin=0.4 * inch,
        title="Md Ashikur Rahman | Industry Resume",
        author="Md Ashikur Rahman",
    )

    story = []
    story.append(Paragraph("MD ASHIKUR RAHMAN", styles["ResumeName"]))
    story.append(Paragraph(_escape(HEADLINE), styles["ResumeHeadline"]))
    story.append(Paragraph(_escape(LOCATION_LINE), styles["ResumeContact"]))
    story.append(Paragraph(contact_pdf_html(), styles["ResumeContact"]))

    story.append(Paragraph("Summary", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(Paragraph(_escape(SUMMARY), styles["Summary"]))

    story.append(Paragraph("Experience", styles["SectionTitle"]))
    story.append(section_rule())
    for title, dates, items in EXPERIENCE:
        block = [entry_header_table(title, dates, styles)]
        block.extend(bullets(items, styles))
        block.append(Spacer(1, 3))
        story.append(KeepTogether(block))

    story.append(Paragraph("Key Projects", styles["SectionTitle"]))
    story.append(section_rule())
    for title, status, items in PROJECTS:
        block = [entry_header_table(title, status, styles)]
        block.extend(bullets(items, styles))
        block.append(Spacer(1, 3))
        story.append(KeepTogether(block))

    story.append(PageBreak())

    pub_block = [
        Paragraph("Peer-Reviewed Publications", styles["SectionTitle"]),
        section_rule(),
    ]
    for pub in PEER_REVIEWED:
        pub_block.append(Paragraph(f"• {_escape(pub)}", styles["PubItem"]))
    story.append(KeepTogether(pub_block))

    preprint_block = [
        Paragraph("Preprints and Manuscripts Under Review", styles["SectionTitle"]),
        section_rule(),
    ]
    for text, label, url in PREPRINTS:
        preprint_block.append(
            Paragraph(
                f'• {_escape(text)} <link href="{url}" color="#9b1b30"><u>{_escape(label)}</u></link>',
                styles["PubItem"],
            )
        )
    story.append(KeepTogether(preprint_block))

    story.append(Paragraph("Education", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(Paragraph("B.Sc. in Computer Science and Engineering", styles["Degree"]))
    story.append(
        Paragraph(
            "American International University-Bangladesh | 2011-2015",
            styles["EntryTitle"],
        )
    )
    story.append(Paragraph("CGPA: 3.87/4.00; WES equivalent: 3.94/4.00", styles["EduMeta"]))
    story.append(
        Paragraph(
            "Magna Cum Laude; Top 3%; Merit Scholarship and Tuition Fee Waiver",
            styles["EduMeta"],
        )
    )

    story.append(Paragraph("Awards", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(
        Paragraph(
            "• Champion, BASIS National ICT Awards 2020 (Retouched.ai); Finalist, APICTA 2021",
            styles["PubItem"],
        )
    )

    story.append(Paragraph("Languages", styles["SectionTitle"]))
    story.append(section_rule())
    story.append(
        Paragraph(
            "• Bengali (Native); English (Professional working proficiency)",
            styles["PubItem"],
        )
    )

    story.append(Paragraph("Skills", styles["SectionTitle"]))
    story.append(section_rule())
    for label, text in SKILLS:
        story.append(
            Paragraph(f"<b>{_escape(label)}</b> {_escape(text)}", styles["SkillLine"])
        )

    doc.build(story)
    print(f"Wrote {PDF_PATH}")


def set_run_font(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E4EAEE")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2, line=1.1)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, color=NAVY_RGB)
    add_bottom_border(p)
    return p


def add_entry_header(doc, title, dates=""):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=1, line=1.1)
    if dates:
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
        run = p.add_run(title)
        set_run_font(run, size=10, bold=True, color=BODY_RGB)
        p.add_run("\t")
        date_run = p.add_run(dates)
        set_run_font(date_run, size=9, bold=False, color=MUTED_RGB)
    else:
        run = p.add_run(title)
        set_run_font(run, size=10, bold=True, color=BODY_RGB)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=1, line=1.15)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=False, color=BODY_RGB)
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "9B1B30")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, before=0, after=2, line=1.0)
    run = name.add_run("MD ASHIKUR RAHMAN")
    set_run_font(run, size=18, bold=True, color=NAVY_RGB)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(headline, before=0, after=2, line=1.1)
    run = headline.add_run(HEADLINE)
    set_run_font(run, size=10, bold=True, color=NAVY_RGB)

    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(location, before=0, after=1, line=1.1)
    run = location.add_run(LOCATION_LINE)
    set_run_font(run, size=9, bold=False, color=NAVY_RGB)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, before=0, after=6, line=1.1)
    for item in CONTACT_PARTS:
        if item[0] == "text":
            run = contact.add_run(item[1])
            set_run_font(run, size=9, bold=False, color=NAVY_RGB)
        else:
            add_hyperlink(contact, item[1], item[2])

    add_section_heading(doc, "Summary")
    summary = doc.add_paragraph()
    set_paragraph_spacing(summary, before=2, after=4, line=1.15)
    summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = summary.add_run(SUMMARY)
    set_run_font(run, size=9.5, color=BODY_RGB)

    add_section_heading(doc, "Experience")
    for title, dates, items in EXPERIENCE:
        add_entry_header(doc, title, dates)
        for item in items:
            add_bullet(doc, item)

    add_section_heading(doc, "Key Projects")
    for title, status, items in PROJECTS:
        add_entry_header(doc, title, status)
        for item in items:
            add_bullet(doc, item)

    doc.add_page_break()

    add_section_heading(doc, "Peer-Reviewed Publications")
    for pub in PEER_REVIEWED:
        add_bullet(doc, pub)

    add_section_heading(doc, "Preprints and Manuscripts Under Review")
    for text, label, url in PREPRINTS:
        p = add_bullet(doc, text + " ")
        add_hyperlink(p, label, url)

    add_section_heading(doc, "Education")
    degree = doc.add_paragraph()
    set_paragraph_spacing(degree, before=2, after=1, line=1.1)
    run = degree.add_run("B.Sc. in Computer Science and Engineering")
    set_run_font(run, size=10, bold=True, color=NAVY_RGB)
    add_entry_header(doc, "American International University-Bangladesh | 2011-2015")
    for line in (
        "CGPA: 3.87/4.00; WES equivalent: 3.94/4.00",
        "Magna Cum Laude; Top 3%; Merit Scholarship and Tuition Fee Waiver",
    ):
        meta = doc.add_paragraph()
        set_paragraph_spacing(meta, before=0, after=1, line=1.15)
        run = meta.add_run(line)
        set_run_font(run, size=9.5, color=BODY_RGB)

    add_section_heading(doc, "Awards")
    add_bullet(
        doc,
        "Champion, BASIS National ICT Awards 2020 (Retouched.ai); Finalist, APICTA 2021",
    )

    add_section_heading(doc, "Languages")
    add_bullet(doc, "Bengali (Native); English (Professional working proficiency)")

    add_section_heading(doc, "Skills")
    for label, text in SKILLS:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=1, line=1.15)
        run = p.add_run(label + " ")
        set_run_font(run, size=9.5, bold=True, color=BODY_RGB)
        run2 = p.add_run(text)
        set_run_font(run2, size=9.5, bold=False, color=BODY_RGB)

    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build_pdf()
    build_docx()
